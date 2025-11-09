# ============================================================================
# DOCX EXTRACTOR - Extract content from Word documents
# ============================================================================

from .base_extractor import BaseExtractor
from typing import Dict, Any, List
import logging

class DOCXExtractor(BaseExtractor):
    """Extract complete content from Word and rich text documents"""
    
    def extract(self) -> Dict[str, Any]:
        """
        Extract all content from Word document (.docx, .doc, .odt, .rtf)
        
        Returns:
            Dictionary with text, tables, images, styles, metadata
        """
        try:
            ext = self.file_path.suffix.lower()
            
            # Only .docx is fully supported by python-docx
            if ext == '.docx':
                return self._extract_docx()
            else:
                # For .doc, .odt, .rtf - try to extract as text
                self.logger.warning(f"{ext} format has limited support - extracting as text only")
                return self._extract_as_text()
                
        except Exception as e:
            self.logger.error(f"Document extraction failed: {e}")
            return self.create_result_dict(
                content={'error': str(e)},
                status='error',
                error_message=str(e)
            )
    
    def _extract_docx(self) -> Dict[str, Any]:
        """Extract full DOCX content"""
        try:
            from docx import Document
            
            doc = Document(self.file_path)
            
            # Extract all components
            content = {
                'text': self._extract_text(doc),
                'paragraphs': self._extract_paragraphs(doc),
                'tables': self._extract_tables(doc),
                'styles': self._extract_styles(doc),
                'sections': len(doc.sections),
                'metadata': self._extract_metadata(doc),
                'extraction_method': 'python-docx'
            }
            
            return self.create_result_dict(content)
            
        except ImportError:
            error_msg = "python-docx not installed. Install: pip install python-docx"
            self.logger.error(error_msg)
            return self.create_result_dict(
                content={'error': error_msg},
                status='error',
                error_message=error_msg
            )
        except Exception as e:
            self.logger.error(f"DOCX extraction failed: {e}")
            return self.create_result_dict(
                content={'error': str(e)},
                status='error',
                error_message=str(e)
            )
    
    def _extract_as_text(self) -> Dict[str, Any]:
        """Fallback text extraction for .doc, .odt, .rtf formats"""
        try:
            # Try to read as plain text with encoding detection
            text_content = self.read_file_content()
            
            content = {
                'text': text_content,
                'note': f'{self.file_path.suffix} format - limited extraction (text only)',
                'extraction_method': 'fallback_text_reader',
                'lines': len(text_content.split('\n')),
                'characters': len(text_content)
            }
            
            return self.create_result_dict(content)
            
        except Exception as e:
            self.logger.error(f"Text fallback extraction failed: {e}")
            return self.create_result_dict(
                content={'error': str(e)},
                status='error',
                error_message=str(e)
            )
    
    def _extract_text(self, doc) -> str:
        """Extract all text maintaining structure"""
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return '\n\n'.join(paragraphs)
    
    def _extract_paragraphs(self, doc) -> List[Dict]:
        """Extract paragraphs with metadata"""
        paragraphs = []
        for idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                paragraphs.append({
                    'index': idx,
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal'
                })
        return paragraphs
    
    def _extract_tables(self, doc) -> List[Dict]:
        """Extract all tables"""
        tables = []
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            
            tables.append({
                'table_index': table_idx + 1,
                'rows': len(table.rows),
                'columns': len(table.columns),
                'data': table_data
            })
        return tables
    
    def _extract_styles(self, doc) -> List[str]:
        """Extract list of styles used"""
        styles = set()
        for para in doc.paragraphs:
            if para.style:
                styles.add(para.style.name)
        return sorted(list(styles))
    
    def _extract_metadata(self, doc) -> Dict:
        """Extract document metadata"""
        core_props = doc.core_properties
        return {
            'title': core_props.title or '',
            'author': core_props.author or '',
            'subject': core_props.subject or '',
            'keywords': core_props.keywords or '',
            'comments': core_props.comments or '',
            'created': core_props.created.isoformat() if core_props.created else '',
            'modified': core_props.modified.isoformat() if core_props.modified else '',
            'last_modified_by': core_props.last_modified_by or '',
            'revision': core_props.revision or 0,
        }
